#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Classification des DOCX (première page + nom de fichier) en EDB / NDC / AUTRES.

Ordre d'évaluation (règles validées) :
  1) NDC si code en première page
  2) EDB si le nom contient "edb" (insensible casse)
  3) EDB si le nom contient "expr + besoin(s)" (Règle A, insensible casse/accents, séparateurs libres, "de" optionnel)
  4) EDB si le nom contient "eb" ET qu'aucun code NDC n'est présent en première page
  5) NDC si un code est détecté dans le nom du fichier
  6) EDB si la première page contient une des phrases EDB (insensible casse/accents)
  7) AUTRES sinon

Caractéristiques :
- NDC multi-clients : CAPS et AVEM (tolérance aux espaces internes entre lettres).
- Extraction "1re page" robuste (paragraphes, tables, textboxes, header/footer) avec namespaces + safe_xpath.
- Si le DOCX est illisible, on CLASSIFIE quand même par le NOM (et on copie).
- Rapport Excel écrit dans le parent de --docx-dir (ex: datas/classify_report.xlsx).

⚠️ Modif demandée ici :
- La partie "année" du motif NDC n’est plus limitée aux chiffres : elle accepte désormais 4 caractères alphanumériques (ex: `A2B3`).
"""

import argparse
from datetime import datetime
from pathlib import Path
import re
import shutil
import unicodedata

import pandas as pd
from docx import Document
from docx.oxml.ns import qn

# ---------- Configuration par défaut ----------
DEFAULT_INPUT_DIR = "docx"
DEFAULT_OUTPUT_DIR = "classified_docx"
DEFAULT_ON_EXISTS = "skip"      # skip | overwrite | suffix
DEFAULT_FIRST_PAGE_CHAR_LIMIT = 12000

# ---------- Namespaces XML pour les XPath ----------
NS = {
    "w":  "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "a":  "http://schemas.openxmlformats.org/drawingml/2006/main",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "v":  "urn:schemas-microsoft-com:vml",
}

def safe_xpath(el, expression):
    """Exécute un XPath en sécurité et renvoie [] en cas d'échec (évite XPathEvalError)."""
    try:
        return el.xpath(expression, namespaces=NS)
    except Exception:
        return []

# ---------- EDB Tokens (insensibles accents/casse) ----------
EDB_TOKENS = [
    "expression de besoin",
    "expression de besoins",
    "expressions de besoins",
]

# RÈGLE A (nom) : expr + (de)? + besoin(s), séparateurs libres (. _ - espaces), accents/casse insensibles
EDB_NAME_ABBR_PATTERN = r"\bexpr(?:ession)?[\W_]*(?:de[\W_]*)?besoin(?:s)?\b"
EDB_NAME_ABBR_REGEX = re.compile(EDB_NAME_ABBR_PATTERN)

# ---------- Regex NDC (large et tolérante) ----------
# Clients acceptés avec tolérance espaces internes : "CAPS" -> C\s*A\s*P\s*S ; "AVEM" -> A\s*V\s*E\s*M
CLIENTS = ["CAPS", "AVEM"]
CLIENT_PATTERNS = [r"\s*".join(list(c)) for c in CLIENTS]   # ["C\s*A\s*P\s*S", "A\s*V\s*E\s*M"]
CLIENT_ALT = "(?:" + "|".join(CLIENT_PATTERNS) + ")"

# Séparateurs libres entre segments : espace, underscore, hyphens (ASCII et typographiques)
SEP = r"[ \-_\u2011\u2012\u2013\u2014]*"   # 0+ pour tolérer client+année collés (CAPS2023-123)

# Modèle NDC : CLIENT SEP YEAR(4 alphanum) SEP CODE
# - YEAR : 4 caractères alphanumériques (⚠️ élargi vs version précédente)
# - CODE : alphanum + sous-segments '-'/'_' (déjà élargi précédemment)
# Pas d'ancre de fin (pour capter des suffixes "_PF", etc.)
NDC_PATTERN = rf"(?i){CLIENT_ALT}{SEP}[A-Za-z0-9]{{4}}{SEP}[A-Za-z0-9][A-Za-z0-9\-_]*"
NDC_REGEX = re.compile(NDC_PATTERN, flags=re.IGNORECASE)

# ---------- Utils accents ----------
def strip_accents(s: str) -> str:
    if s is None:
        return ""
    return "".join(c for c in unicodedata.normalize("NFD", s) if not unicodedata.combining(c))

# ---------- Extraction 1re page ----------
def paragraph_has_page_break(p) -> bool:
    for br in safe_xpath(p._element, ".//w:br"):
        br_type = br.get(qn("w:type"))
        if (br_type or "").lower() == "page":
            return True
    return False

def element_text_runs(el) -> str:
    texts = []
    for t in safe_xpath(el, ".//w:t"):
        if t.text:
            texts.append(t.text)
    return "\n".join(texts)

def extract_header_footer_text(doc: Document) -> str:
    parts = []
    try:
        sec = doc.sections[0]
        if getattr(sec, "header", None):
            for p in sec.header.paragraphs:
                if p.text:
                    parts.append(p.text)
            for tbl in safe_xpath(sec.header._element, ".//w:tbl"):
                parts.append(element_text_runs(tbl))
        if getattr(sec, "footer", None):
            for p in sec.footer.paragraphs:
                if p.text:
                    parts.append(p.text)
            for tbl in safe_xpath(sec.footer._element, ".//w:tbl"):
                parts.append(element_text_runs(tbl))
    except Exception:
        pass
    return "\n".join(filter(None, parts))

def extract_first_page_text(docx_path: Path, char_limit: int) -> str:
    """
    "Approx first page" robuste :
      - header/footer section 1
      - corps du document : paragraphs + tables + textboxes
      - stop au 1er saut de page ou à char_limit
    """
    doc = Document(str(docx_path))
    parts = []

    hf = extract_header_footer_text(doc)
    if hf:
        parts.append(hf)

    total_len = sum(len(p) + 1 for p in parts)
    body = doc.element.body

    for child in body.iterchildren():
        tag = child.tag.rsplit("}", 1)[-1].lower()

        if tag == "p":
            p_txt = element_text_runs(child)
            if p_txt:
                parts.append(p_txt)
                total_len += len(p_txt) + 1
            for br in safe_xpath(child, ".//w:br"):
                br_type = br.get(qn("w:type"))
                if (br_type or "").lower() == "page":
                    return "\n".join(parts)

        elif tag == "tbl":
            t_txt = element_text_runs(child)
            if t_txt:
                parts.append(t_txt)
                total_len += len(t_txt) + 1

        # Textboxes (contenu texte encapsulé)
        txbx_chunks = []
        for txbx in safe_xpath(child, ".//w:txbxContent"):
            txbx_chunks.append(element_text_runs(txbx))
        if txbx_chunks:
            tx = "\n".join(filter(None, txbx_chunks))
            if tx:
                parts.append(tx)
                total_len += len(tx) + 1

        if total_len >= char_limit:
            break

    return "\n".join(parts)[:char_limit]

# ---------- Copies ----------
def ensure_dirs(base_out: Path):
    (base_out / "edb").mkdir(parents=True, exist_ok=True)
    (base_out / "ndc").mkdir(parents=True, exist_ok=True)
    (base_out / "autres").mkdir(parents=True, exist_ok=True)

def safe_copy(src: Path, dst_dir: Path, on_exists: str):
    dst_dir.mkdir(parents=True, exist_ok=True)
    dst = dst_dir / src.name
    if dst.exists():
        if on_exists == "skip":
            return dst, "skipped_existing"
        elif on_exists == "overwrite":
            shutil.copy2(src, dst)
            return dst, "overwritten"
        elif on_exists == "suffix":
            stem, ext = src.stem, src.suffix
            ts = datetime.now().strftime("%Y%m%d_%H%M%S")
            suffixed = dst_dir / f"{stem}_{ts}{ext}"
            shutil.copy2(src, suffixed)
            return suffixed, "copied_with_suffix"
        else:
            raise ValueError(f"on_exists invalide: {on_exists}")
    else:
        shutil.copy2(src, dst)
        return dst, "copied"

# ---------- Détections ----------
def detect_ndc_in_first_page(text: str) -> tuple[bool, str]:
    m = NDC_REGEX.search(text or "")
    if m:
        return True, f"pattern:{m.group(0)} source:first_page"
    return False, ""

def detect_ndc_in_filename(filename: str) -> tuple[bool, str]:
    m = NDC_REGEX.search(filename or "")
    if m:
        return True, f"pattern:{m.group(0)} source:filename"
    return False, ""

def detect_edb_in_first_page(text: str) -> tuple[bool, str]:
    norm_text = strip_accents(text).lower()
    for token in EDB_TOKENS:
        if token in norm_text:
            return True, f"contains_first_page:'{token}'"
    return False, ""

def detect_edb_phrases_in_filename(filename: str) -> tuple[bool, str]:
    """EDB si le nom contient une des phrases EDB (insensible casse/accents)."""
    norm_name = strip_accents(filename).lower()
    for token in EDB_TOKENS:
        if token in norm_name:
            return True, f"filename_contains_phrase:'{token}'"
    return False, ""

def detect_edb_abbrev_in_filename(filename: str) -> tuple[bool, str]:
    """
    RÈGLE A (prudente) : EDB si le nom présente 'expr' (+ 'ession' optionnel) puis (optionnel 'de') puis 'besoin(s)',
    avec séparateurs libres (. _ - espaces).
    """
    norm_name = strip_accents(filename).lower()
    if EDB_NAME_ABBR_REGEX.search(norm_name):
        return True, "filename_contains_abbrev:'expr...besoin(s)'"
    return False, ""

# ---------- Classification ----------
def classify(first_page_text: str, filename: str, content_read_ok: bool) -> tuple[str, str]:
    """
    Respecte l'ordre demandé, avec fallback par nom si contenu illisible.
    """
    filename_lower = (filename or "").lower()

    # 1) NDC si code en première page (si lisible)
    if content_read_ok:
        ndc_first, reason_ndc_first = detect_ndc_in_first_page(first_page_text)
        if ndc_first:
            return "NDC", reason_ndc_first

    # 2) EDB si nom contient 'edb'
    if "edb" in filename_lower:
        return "EDB", "filename_contains:edb"

    # 3) EDB si nom contient une des phrases EDB OU l'abréviation 'expr...besoin(s)'
    edb_name_phrase, reason_phrase = detect_edb_phrases_in_filename(filename)
    if edb_name_phrase:
        return "EDB", reason_phrase
    edb_name_abbrev, reason_abbrev = detect_edb_abbrev_in_filename(filename)
    if edb_name_abbrev:
        return "EDB", reason_abbrev

    # 4) EDB si nom contient 'eb' ET pas de code NDC en 1re page
    if "eb" in filename_lower:
        if content_read_ok:
            ndc_first, _ = detect_ndc_in_first_page(first_page_text)
            if not ndc_first:
                return "EDB", "filename_contains:eb AND no_ndc_on_first_page"
        else:
            return "EDB", "filename_contains:eb AND content_unreadable"

    # 5) NDC si code dans le nom
    ndc_name, reason_ndc_name = detect_ndc_in_filename(filename)
    if ndc_name:
        return "NDC", reason_ndc_name

    # 6) EDB si tokens EDB dans 1re page (si lisible)
    if content_read_ok:
        edb_text, reason_edb = detect_edb_in_first_page(first_page_text)
        if edb_text:
            return "EDB", reason_edb

    # 7) AUTRES
    return "AUTRES", ""

# ---------- Main ----------
def main():
    parser = argparse.ArgumentParser(description="Classement DOCX en EDB / NDC / AUTRES (1ère page + nom).")
    parser.add_argument("--docx-dir", default=DEFAULT_INPUT_DIR,
                        help="Dossier d'entrée contenant les .docx (défaut: docx)")
    parser.add_argument("--output-dir", default=DEFAULT_OUTPUT_DIR,
                        help="Dossier racine de sortie (défaut: classified_docx)")
    parser.add_argument("--on-exists", choices=["skip", "overwrite", "suffix"], default=DEFAULT_ON_EXISTS,
                        help="Politique en cas de collision de nom (défaut: skip)")
    parser.add_argument("--recursive", action="store_true",
                        help="Parcourir récursivement le dossier d'entrée")
    parser.add_argument("--first-page-char-limit", type=int, default=DEFAULT_FIRST_PAGE_CHAR_LIMIT,
                        help="Troncature si pas de saut de page explicite (défaut: 12000)")
    parser.add_argument("--debug-first-pages", action="store_true",
                        help="Sauvegarde le texte extrait (approx. 1ère page) dans classified_docx/_debug_first_pages")
    args = parser.parse_args()

    in_dir = Path(args.docx_dir).resolve()
    base_out = Path(args.output_dir).resolve()
    ensure_dirs(base_out)

    debug_dir = base_out / "_debug_first_pages"
    if args.debug_first_pages:
        debug_dir.mkdir(parents=True, exist_ok=True)

    candidates = list(in_dir.rglob("*.docx")) if args.recursive else list(in_dir.glob("*.docx"))

    records = []
    total = len(candidates)
    print(f"[INFO] {total} fichier(s) .docx à traiter dans: {in_dir}")

    for i, path in enumerate(sorted(candidates, key=lambda p: str(p).lower()), start=1):
        try:
            rel = path.relative_to(in_dir)
        except Exception:
            rel = path.name

        classification = "ERREUR"
        reason = ""
        dest_path = None
        copy_status = "not_copied"

        first_page = ""
        content_read_ok = True

        # Lecture + extraction "1re page"
        try:
            first_page = extract_first_page_text(path, char_limit=args.first_page_char_limit)
            if args.debug_first_pages:
                try:
                    with open(debug_dir / f"{path.stem}.txt", "w", encoding="utf-8") as fdbg:
                        fdbg.write(first_page)
                except Exception:
                    pass
        except Exception as e:
            # On continue malgré tout : fallback par nom
            content_read_ok = False
            reason = f"content_unreadable:{type(e).__name__}"

        try:
            # Classification
            cls, rsn = classify(first_page, path.name, content_read_ok)
            if (not content_read_ok) and rsn:
                rsn = rsn + " | " + reason
            elif (not content_read_ok) and (not rsn):
                rsn = reason or "content_unreadable"

            classification, reason = cls, rsn

            # Dossier cible
            if classification == "EDB":
                target_dir = base_out / "edb"
            elif classification == "NDC":
                target_dir = base_out / "ndc"
            elif classification == "AUTRES":
                target_dir = base_out / "autres"
            else:
                target_dir = None

            if target_dir is not None:
                dest_path, copy_status = safe_copy(path, target_dir, args.on_exists)

            print(f"[{i}/{total}] {rel} -> {classification} ({reason or 'no_reason'})")

        except Exception as e:
            classification = "ERREUR"
            reason = f"exception:{type(e).__name__}: {e}"
            print(f"[{i}/{total}] {rel} -> ERREUR ({reason})")

        records.append({
            "filename": path.name,
            "original_path": str(path),
            "classification": classification,
            "reason": reason,
            "destination_path": "" if dest_path is None else str(dest_path),
            "copy_status": copy_status,
        })

    # Rapport Excel -> parent de --docx-dir (ex: datas/classify_report.xlsx)
    repo_root = in_dir.parent
    report_path = repo_root / "classify_report.xlsx"
    df = pd.DataFrame.from_records(
        records,
        columns=["filename", "original_path", "classification", "reason", "destination_path", "copy_status"],
    )
    df.to_excel(report_path, index=False)
    print(f"[OK] Rapport écrit : {report_path}")
    print("[OK] Terminé.")


if __name__ == "__main__":
    main()
