#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Conversion des fichiers DOCX classés (NDC/EDB) en Markdown.

Source (lecture) :
  classified_docx/
    ├─ ndc/
    └─ edb/

Cible (écriture) :
  markdown/
    ├─ ndc/
    └─ edb/

Fonctionnement :
- Parcours récursif des répertoires source
- Conversion .docx -> .md (titres, gras, italique, listes, tableaux, paragraphes)
- Politique en cas de collision : skip | overwrite | suffix

Dépendances :
- python-docx
- (standard library) pandas non requis ici

Limitations :
- Mise en forme avancée (numérotation hiérarchique, images, notes de bas de page)
  non rendue ; les images sont indiquées sous forme de commentaires Markdown.
"""

import argparse
from pathlib import Path
from datetime import datetime
import unicodedata
import re
import sys

from docx import Document


# -------------------------
# Helpers système / FS
# -------------------------
def ensure_dirs(*paths: Path):
    for p in paths:
        p.mkdir(parents=True, exist_ok=True)


def resolve_collision(dst_path: Path, on_exists: str) -> Path:
    if not dst_path.exists():
        return dst_path
    if on_exists == "skip":
        return dst_path  # l'appelant testera l'existence et skippera
    if on_exists == "overwrite":
        return dst_path
    if on_exists == "suffix":
        stem, ext = dst_path.stem, dst_path.suffix
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        return dst_path.with_name(f"{stem}_{ts}{ext}")
    raise ValueError(f"on_exists invalide: {on_exists}")


# -------------------------
# Helpers Markdown
# -------------------------
def strip_accents(s: str) -> str:
    if s is None:
        return ""
    return "".join(c for c in unicodedata.normalize("NFD", s) if not unicodedata.combining(c))


def runs_to_markdown(paragraph) -> str:
    """
    Convertit les "runs" d'un paragraphe en Markdown basique (gras/italique/code).
    - **gras**, *italique*, `code` (si style "Code")
    - concaténation textuelle sinon
    """
    parts = []
    for run in paragraph.runs:
        text = run.text or ""
        if not text:
            continue

        # rudimentaire : si la police/format suggère du code (style nommé "Code")
        is_code = (paragraph.style and paragraph.style.name and "code" in paragraph.style.name.lower())
        bold = run.bold
        italic = run.italic

        if is_code:
            parts.append(f"`{text}`")
        else:
            if bold and italic:
                parts.append(f"***{text}***")
            elif bold:
                parts.append(f"**{text}**")
            elif italic:
                parts.append(f"*{text}*")
            else:
                parts.append(text)
    return "".join(parts).strip()


def para_to_markdown(p) -> str:
    """
    Conversion d'un paragraphe en Markdown :
    - Titres : styles "Heading 1..6" -> # .. ######
    - Listes : style "List Paragraph" -> "- ..."
    - Sinon : paragraphe standard
    """
    style_name = (p.style.name if p.style else "") or ""
    text_md = runs_to_markdown(p)
    if not text_md.strip():
        return ""

    # Titres (Heading n)
    m = re.match(r"Heading\s*(\d+)", style_name, flags=re.I)
    if m:
        level = max(1, min(6, int(m.group(1))))
        return f"{'#' * level} {text_md}"

    # Liste simple
    if "list" in style_name.lower():
        return f"- {text_md}"

    return text_md


def table_to_markdown(table) -> str:
    """
    Conversion d'un tableau docx en table Markdown simple.
    - Première ligne utilisée comme en-tête
    - Alignement par défaut : gauche
    """
    rows = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            # concatène le texte de chaque paragraphe de la cellule
            cell_text = "\n".join((runs_to_markdown(p) for p in cell.paragraphs)).strip()
            # échappe les pipe
            cell_text = cell_text.replace("|", r"\|")
            cells.append(cell_text)
        rows.append(cells)

    if not rows:
        return ""

    header = rows[0]
    align = ["---"] * len(header)
    md = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(align) + " |",
    ]
    for r in rows[1:]:
        md.append("| " + " | ".join(r) + " |")
    return "\n".join(md)


def docx_to_markdown(docx_path: Path) -> str:
    """
    Convertit un document DOCX en Markdown basique.
    - parcourt les paragraphes et tables dans l'ordre
    - insère des lignes vides entre blocs pour lisibilité
    - images : insère un commentaire markdown indiquant la présence
    """
    doc = Document(str(docx_path))
    out_lines = []

    # En-tête optionnelle : titre avec nom de fichier
    out_lines.append(f"<!-- Source: {docx_path.name} -->")
    out_lines.append(f"# {docx_path.stem}")
    out_lines.append("")

    # Approche simple et fiable : paragraphs puis tables
    # (la majorité des documents ne nécessitent pas l'ordre intercalé strict)
    for p in doc.paragraphs:
        line = para_to_markdown(p)
        if line:
            out_lines.append(line)

    # Tables
    if doc.tables:
        out_lines.append("")
        out_lines.append("> **Tableaux**")
        out_lines.append("")
        for t in doc.tables:
            md_table = table_to_markdown(t)
            if md_table:
                out_lines.append(md_table)
                out_lines.append("")

    # Images (si présentes) : docx stocke images dans relationships ; on ajoute un rappel
    rels = getattr(doc.part, "rels", {})
    has_images = any(
        (getattr(r._target, "content_type", "") or "").startswith("image/") for r in rels.values()
    )
    if has_images:
        out_lines.append("> _Ce document contient des images non extraites dans ce rendu Markdown._")

    # Normalisation : supprimer doublons de lignes vides
    cleaned = []
    last_blank = False
    for line in out_lines:
        if line.strip() == "":
            if not last_blank:
                cleaned.append("")
            last_blank = True
        else:
            cleaned.append(line)
            last_blank = False

    return "\n".join(cleaned).rstrip() + "\n"


# -------------------------
# Programme principal
# -------------------------
def main():
    parser = argparse.ArgumentParser(description="Convertit les DOCX classés (NDC/EDB) en Markdown.")
    parser.add_argument("--classified-dir", default="classified_docx",
                        help="Dossier racine des DOCX classés (défaut: classified_docx)")
    parser.add_argument("--markdown-dir", default="markdown",
                        help="Dossier racine de sortie Markdown (défaut: markdown)")
    parser.add_argument("--on-exists", choices=["skip", "overwrite", "suffix"], default="skip",
                        help="Politique en cas de collision de nom (défaut: skip)")
    parser.add_argument("--recursive", action="store_true",
                        help="Parcourt récursivement ndc/ et edb/")
    args = parser.parse_args()

    classified_root = Path(args.classified_dir).resolve()
    src_ndc = classified_root / "ndc"
    src_edb = classified_root / "edb"

    md_root = Path(args.markdown_dir).resolve()
    dst_ndc = md_root / "ndc"
    dst_edb = md_root / "edb"
    ensure_dirs(dst_ndc, dst_edb)

    if not src_ndc.exists() and not src_edb.exists():
        print(f"[ERREUR] Répertoires source introuvables sous {classified_root} (ndc/ et edb/).", file=sys.stderr)
        sys.exit(1)

    def iter_docx(src_dir: Path):
        if not src_dir.exists():
            return []
        if args.recursive:
            return list(src_dir.rglob("*.docx"))
        return list(src_dir.glob("*.docx"))

    files = [(p, "ndc") for p in iter_docx(src_ndc)] + [(p, "edb") for p in iter_docx(src_edb)]
    print(f"[INFO] {len(files)} fichier(s) .docx détecté(s) dans {classified_root}/(ndc|edb)")

    converted = 0
    skipped = 0
    errors = 0

    for i, (path, category) in enumerate(sorted(files, key=lambda t: str(t[0]).lower()), start=1):
        print(f"[{i}/{len(files)}] {category.upper()} : {path.name}")
        try:
            md_text = docx_to_markdown(path)
            out_dir = dst_ndc if category == "ndc" else dst_edb
            out_path = out_dir / (path.stem + ".md")

            # Collisions
            final_path = resolve_collision(out_path, args.on_exists)
            if out_path.exists() and args.on_exists == "skip":
                print(f"  - Skip (existe déjà) → {out_path}")
                skipped += 1
                continue

            out_dir.mkdir(parents=True, exist_ok=True)
            out_path_to_write = final_path

            with open(out_path_to_write, "w", encoding="utf-8") as f:
                f.write(md_text)

            if out_path_to_write != out_path:
                print(f"  - Écrit (suffix) → {out_path_to_write}")
            else:
                print(f"  - Écrit → {out_path_to_write}")
            converted += 1

        except Exception as e:
            print(f"  ! ERREUR : {type(e).__name__}: {e}", file=sys.stderr)
            errors += 1

    print("")
    print(f"[RÉSUMÉ] Convertis: {converted} • Ignorés (skip): {skipped} • Erreurs: {errors}")
    print(f"[SORTIE ] Markdown : {md_root}")


if __name__ == "__main__":
    main()
